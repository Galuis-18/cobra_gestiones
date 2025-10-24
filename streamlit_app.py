import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
import io 
import tempfile
import os
from docx import Document
from docx.shared import Inches  
import warnings
import subprocess
#from docx2pdf import convert

## Funciones viva españa
def fetch_date(texto):
    return str(texto)[:7] + "-" + str(texto)[7:9]

def fetch_time(texto):
    return str(texto)[10:18]

def get_mode(x):
    m = x.mode()
    return m.iloc[0] if not m.empty else np.nan


def generate_report(uploaded_file):
    """
    Descripción:
    """
    
    # Ignorar advertencias 
    warnings.filterwarnings('ignore', category=UserWarning)

    ## Leer archivo cargado de excel
    try:
        gestiones = pd.read_excel(uploaded_file)
    except Exception as e:
        raise Exception(f"No se pudo leer el archivo Excel. Error: {e}")

    if gestiones.columns[0] == 'Gestiones desde APP':
        gestiones.columns = gestiones.iloc[0]
        gestiones = gestiones.iloc[1:]
    
    # Validar columnas necesarias antes de procesar
    required_cols = ['Fecha', 'No. de Cobrador', 'No. de Contrato', 'Monto']
    if not all(col in gestiones.columns for col in required_cols):
        missing = [col for col in required_cols if col not in gestiones.columns]
        raise Exception(f"El archivo no es válido. Faltan las siguientes columnas: {', '.join(missing)}")

    gestiones['dia'] = gestiones['Fecha'].apply(fetch_date)
    gestiones['time'] = gestiones['Fecha'].apply(fetch_time)
    gestiones = gestiones.rename(columns={"No. de Cobrador": "empleado", "No. de Contrato": "contrato"})
    df = gestiones.copy()

    # Quitar aquellos que en el nombre de empleado tengan la palabra prueba
    df = df[~df['empleado'].str.contains('prueba', case=False)]

    # 2. Procesamiento de Fechas y Tiempos
    try:
        df['datetime_completo'] = pd.to_datetime(df['dia'] + ' ' + df['time'])
        df['dia_dt'] = pd.to_datetime(df['dia'])
        df = df.sort_values(by=['empleado', 'datetime_completo'])
        df['tiempo_entre_gestiones'] = df.groupby(['empleado', 'dia_dt'])['datetime_completo'].diff()
        df['segundos_entre_gestiones'] = df['tiempo_entre_gestiones'].dt.total_seconds()
    except Exception as e:
        # Lanzamos una excepción para que el UI de Streamlit la capture
        raise Exception(f"Asegurate que la columna fecha esté en formato: 2025-1001_07:44:51_O005587. Detalle: {e}")

    # Creación del word
    doc = Document()
    doc.add_heading('Análisis por Cobrador', 0)
    doc.add_paragraph('En este reporte se hace un análisis de las medidas de tendencia y algunas estadísticas de los tiempos de gestión por cobrador.')
    doc.add_paragraph('Se recomienda altamente tomar como medida principal para el análisis la mediana. Esto ya que el promedio es una variable susceptible a los valores extremos.')
    doc.add_paragraph('Si a lo largo de cada día hay una gestión que dure más de una hora, eso terminará afectando el promedio general a largo plazo.')
    doc.add_paragraph('De la misma forma la moda puede estar fallando por múltiples gestiones cortas, es decir, si encuentra dos valores repetidos y las demás gestiones no duraron lo mismo, estos serán tomados como moda general.')
    doc.add_paragraph('Reporte con estadísticas de ventas y análisis de tiempos por cobrador.')

    empleados_unicos = df['empleado'].unique()
    total_empleados = len(empleados_unicos)

    if total_empleados == 0:
        doc.add_paragraph("No se encontraron empleados en el archivo después de filtrar la palabra 'prueba'.")
        # Guardar en memoria y devolver
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()

    # --- 4. BUCLE POR CADA EMPLEADO ---
    for i, empleado_id in enumerate(empleados_unicos):
        
        # Filtrar datos del empleado
        df_empleado = df[df['empleado'] == empleado_id].copy()
        tiempos_empleado = df_empleado['segundos_entre_gestiones'].dropna()

        if tiempos_empleado.empty:
            continue

        # Cálculo de Estadísticas
        total_gestiones_periodo = len(df_empleado)
        
        try:
            moda_general_seg = tiempos_empleado.mode().iloc[0]
            moda_general_str = f"{moda_general_seg:.2f} segundos"
        except IndexError:
            moda_general_str = "N/A"

        monto_total_acumulado = df_empleado['Monto'].sum()
        montos_diarios = df_empleado.groupby('dia_dt')['Monto'].sum()
        monto_promedio_dia = montos_diarios.mean()
        gestiones_por_dia = df_empleado.groupby('dia_dt').size()
        gestiones_promedio_dia = gestiones_por_dia.mean()

        # Cálculo de Datos para Gráfica de Evolución
        daily_stats = df_empleado.groupby('dia_dt')['segundos_entre_gestiones'].agg(
            mediana_diaria='median',
            moda_diaria=get_mode
        ).dropna()
        daily_stats = daily_stats.sort_index()

        # Crear Gráfica 1: Histograma (General)
        mem_fig_hist = io.BytesIO()
        plt.figure(figsize=(8, 4))
        tiempos_filtrados_hist = tiempos_empleado[tiempos_empleado < 3600]
        num_gestiones_no_pasadas = len(tiempos_empleado) - len(tiempos_filtrados_hist)

        if not tiempos_filtrados_hist.empty:
            sns.histplot(tiempos_filtrados_hist, bins=50, kde=True)
        else:
            sns.histplot(tiempos_empleado, bins=50, kde=True)

        plt.title(f'Distribución de Tiempos - {empleado_id}')
        plt.xlabel('Segundos entre Gestiones (Filtrado a 1 hora)')
        plt.ylabel('Frecuencia')
        plt.savefig(mem_fig_hist, format='png', bbox_inches='tight')
        plt.close()

        # Crear Gráfica 2: Serie de Tiempo (Evolución Diaria)
        mem_fig_ts = None
        if len(daily_stats) > 1:
            mem_fig_ts = io.BytesIO()
            plt.figure(figsize=(9, 4.5))
            plt.plot(daily_stats.index, daily_stats['mediana_diaria'], label='Mediana Diaria', marker='o', linestyle='--')
            plt.plot(daily_stats.index, daily_stats['moda_diaria'], label='Moda Diaria', marker='x', linestyle='-')
            plt.title(f'Evolución de Tiempos Diarios - {empleado_id}')
            plt.xlabel('Fecha')
            plt.ylabel('Segundos')
            plt.legend()
            plt.grid(True, which='both', linestyle=':', linewidth=0.5)
            plt.gcf().autofmt_xdate()
            plt.tight_layout()
            plt.savefig(mem_fig_ts, format='png', bbox_inches='tight')
            plt.close()

        # Estadísticas
        doc.add_heading(f'Empleado: {empleado_id}', level=3)
        doc.add_paragraph(f"\tMonto total acumulado: ${monto_total_acumulado:,.2f}")
        doc.add_paragraph(f"\tMonto promedio por día: ${monto_promedio_dia:,.2f}")
        doc.add_paragraph(f"\tGestiones promedio por día: {gestiones_promedio_dia:.2f}")
        doc.add_paragraph(f"\tTotal de gestiones en el periodo: {total_gestiones_periodo}")
        doc.add_paragraph(f"\tTotal de gestiones descartadas (arriba de 1 hora): {num_gestiones_no_pasadas}")
        doc.add_paragraph(f"\tTiempo de moda: {moda_general_str}")
        doc.add_paragraph(f"\tMediana general de tiempo por usuario: {tiempos_empleado.median():,.2f}")
        doc.add_paragraph(f"\tPromedio general de tiempo por usuario: {tiempos_empleado.mean():,.2f}")

        doc.add_heading('Distribución de tiempos entre gestiones', level=4)
        doc.add_paragraph('Muestra la frecuencia de los tiempos entre gestiones. Un pico alto a la izquierda significa muchas gestiones rápidas.')
        mem_fig_hist.seek(0) # Regresar al inicio del buffer de imagen
        doc.add_picture(mem_fig_hist, width=Inches(6.0))

        doc.add_heading('Evolución de Tiempos Diarios (Mediana y Moda)', level=4)
        if mem_fig_ts:
            doc.add_paragraph('Muestra cómo cambian los tiempos "típicos" (mediana) y "más frecuentes" (moda) cada día.')
            mem_fig_ts.seek(0) # Regresar al inicio del buffer de imagen
            doc.add_picture(mem_fig_ts, width=Inches(6.0))
        else:
            doc.add_paragraph("No hay suficientes gestiones para poder analizar las tendencias en las gestiones.")

        doc.add_page_break()

    # Guardar el doc
    
    # Crear un buffer de bytes en memoria
    doc_buffer = io.BytesIO()
    # Guardar el documento en ese buffer
    doc.save(doc_buffer)
    # Regresar al inicio del buffer
    doc_buffer.seek(0)
    
    # Devolver los bytes del archivo
    return doc_buffer.getvalue()

# --- 3. INTERFAZ DE USUARIO DE STREAMLIT ---
# Widget para subir el archivo
st.set_page_config(page_title="Generador de Reportes", layout="centered")
st.title("Generador de Reporte de Gestiones 📄")
st.write("Sube el archivo Excel de 'Gestiones APP' para generar el análisis de tiempos y montos por cobrador.")

# Widget para subir el archivo
uploaded_file = st.file_uploader("Selecciona tu archivo .xlsx", type=["xlsx"])

# --- NUEVA LÓGICA DE INTERFAZ ---
if uploaded_file is not None:
    
    # 1. Opción para elegir el formato
    st.write("") # Espacio
    format_choice = st.radio(
        "Elige el formato de descarga:",
        ("Word (.docx)", "PDF (.pdf)"),
        horizontal=True
    )
    st.write("") # Espacio
    
    # 2. Botón para generar (o se podría hacer automático)
    if st.button("Generar Reporte 🚀"):
        
        with st.spinner("Procesando archivo... Esto puede tardar unos segundos... ⏳"):
            try:
                # 3. Siempre generamos el .docx primero
                docx_bytes = generate_report(uploaded_file)
                st.success("¡Reporte procesado con éxito! 🎉")

                # 4. Lógica de descarga
                if format_choice == "Word (.docx)":
                    st.download_button(
                        label="Descargar Reporte (.docx) 📥",
                        data=docx_bytes,
                        file_name="Reporte_Tiempos_Cobradores.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                elif format_choice == "PDF (.pdf)":
                    with st.spinner("Convirtiendo a PDF... (Esto puede tardar un poco más) 🔄"):
                        
                        pdf_bytes = None
                        temp_docx_path = None
                        temp_pdf_dir = None

                        try:
                            # 1. Crear un directorio temporal para la salida
                            with tempfile.TemporaryDirectory() as temp_pdf_dir:
                                
                                # 2. Crear un archivo .docx temporal
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
                                    temp_docx.write(docx_bytes)
                                    temp_docx_path = temp_docx.name

                                # 3. Construir el comando de LibreOffice
                                # Esto le dice a libreoffice que convierta el docx a pdf
                                # y guarde el resultado en el directorio temporal
                                command = [
                                    "libreoffice",
                                    "--headless",
                                    "--convert-to", "pdf",
                                    "--outdir", temp_pdf_dir,
                                    temp_docx_path
                                ]
                                
                                # 4. Ejecutar el comando
                                result = subprocess.run(command, capture_output=True, text=True, timeout=30)

                                # 5. Revisar si hubo errores
                                if result.returncode != 0:
                                    st.error("Error durante la conversión con LibreOffice:")
                                    st.code(result.stderr) # Muestra el error
                                    raise Exception(f"Fallo de LibreOffice: {result.stderr}")

                                # 6. Encontrar el archivo PDF de salida
                                # El PDF tendrá el mismo nombre que el .docx
                                pdf_filename = os.path.basename(temp_docx_path).replace(".docx", ".pdf")
                                temp_pdf_path = os.path.join(temp_pdf_dir, pdf_filename)

                                if not os.path.exists(temp_pdf_path):
                                    raise Exception("El archivo PDF no fue generado por LibreOffice.")

                                # 7. Leer los bytes del PDF generado
                                with open(temp_pdf_path, "rb") as f:
                                    pdf_bytes = f.read()

                            # 8. Mostrar el botón de descarga
                            st.download_button(
                                label="Descargar Reporte (.pdf) 📥",
                                data=pdf_bytes,
                                file_name="Reporte_Tiempos_Cobradores.pdf",
                                mime="application/pdf"
                            )

                        except Exception as convert_error:
                            st.error(f"Error durante la conversión a PDF: {convert_error}")
                            st.info("La conversión a PDF falló, pero puedes descargar la versión .docx:")
                            # Fallback al botón de .docx
                            st.download_button(
                                label="Descargar Reporte (.docx) 📥",
                                data=docx_bytes,
                                file_name="Reporte_Tiempos_Cobradores.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        finally:
                            # 9. Asegurarse de limpiar el archivo .docx
                            if temp_docx_path and os.path.exists(temp_docx_path):
                                os.remove(temp_docx_path)

            except Exception as e:
                st.error(f"❌ Ocurrió un error al procesar el archivo:")
                st.error(e)
                st.warning("Por favor, verifica el formato del archivo y las columnas requeridas.")
